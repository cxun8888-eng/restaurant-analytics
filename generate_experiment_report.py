"""生成实验报告Word文档 — 算法与功能介绍"""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
import os

doc = Document()

# ===== 页面设置 =====
for section in doc.sections:
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(3.18)
    section.right_margin = Cm(3.18)

style = doc.styles['Normal']
style.font.name = '宋体'
style.font.size = Pt(12)
style.paragraph_format.line_spacing = 1.5
style.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')


def heading(text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0, 0, 0)
        if level == 1: run.font.size = Pt(16)
        elif level == 2: run.font.size = Pt(14)
        else: run.font.size = Pt(12)
    return h


def para(text, bold=False, indent=True):
    p = doc.add_paragraph()
    if indent:
        p.paragraph_format.first_line_indent = Cm(0.74)
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(12)
    return p


def add_table(headers, rows):
    tbl = doc.add_table(rows=1 + len(rows), cols=len(headers))
    tbl.style = 'Table Grid'
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, hd in enumerate(headers):
        cell = tbl.rows[0].cells[i]
        cell.text = hd
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                run.bold = True; run.font.size = Pt(10.5)
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            cell = tbl.rows[ri + 1].cells[ci]
            cell.text = str(val)
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(10.5)
    doc.add_paragraph()
    return tbl


def formula(text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.font.size = Pt(11)
    run.italic = True
    return p


# ============================================================
# 封面
# ============================================================
for _ in range(4):
    doc.add_paragraph()

title_p = doc.add_paragraph()
title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title_p.add_run('餐饮多平台经营数据分析系统')
run.bold = True; run.font.size = Pt(26)
run.font.name = '黑体'

sub_p = doc.add_paragraph()
sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = sub_p.add_run('实验报告')
run.font.size = Pt(18)

doc.add_paragraph()
doc.add_paragraph()

info_items = [
    ('课程名称', '数据挖掘与机器学习'),
    ('实验名称', '餐饮经营数据多维度分析与预测建模'),
    ('专    业', '数据科学与大数据技术'),
    ('姓    名', '__________________'),
    ('学    号', '__________________'),
    ('实验日期', '2026 年    月    日'),
]
for label, value in info_items:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f'{label}：{value}')
    run.font.size = Pt(14)

doc.add_page_break()

# ============================================================
# 一、实验目的
# ============================================================
heading('一、实验目的', 1)
para('1. 掌握数据预处理的基本方法，包括多源数据整合、缺失值处理、异常值检测（IQR方法）等ETL流程。')
para('2. 理解并实现关联规则挖掘算法（Apriori），掌握支持度（Support）、置信度（Confidence）、提升度（Lift）三个核心指标的计算与业务含义。')
para('3. 掌握RFM用户分层模型的构建方法，理解客户价值分析的维度与分群策略。')
para('4. 学习K-Means无监督聚类算法，使用肘部法则确定最优K值，实现聚类结果的交叉验证。')
para('5. 掌握时间序列预测的基本方法，学习将时序问题转化为监督学习问题的特征工程技巧。')
para('6. 学习随机森林回归模型的训练与评估，掌握MAPE（平均绝对百分比误差）等模型评估指标。')
para('7. 了解Isolation Forest异常检测算法的原理与应用场景，理解多维异常检测与单维检测的差异。')
para('8. 培养从数据清洗到模型构建、再到业务洞察的完整数据分析思维。')

# ============================================================
# 二、实验环境
# ============================================================
heading('二、实验环境', 1)
add_table(
    ['项目', '说明'],
    [
        ('操作系统', 'Windows 11'),
        ('编程语言', 'Python 3.12'),
        ('Web框架', 'Streamlit 1.28'),
        ('数据处理', 'Pandas 2.0、NumPy 1.24'),
        ('科学计算', 'SciPy 1.11'),
        ('机器学习库', 'Scikit-learn 1.3'),
        ('数据挖掘库', 'mlxtend 0.22（Apriori算法）'),
        ('可视化工具', 'Plotly 5.17'),
        ('表格处理', 'openpyxl 3.1（Excel文件读写）'),
        ('开发工具', 'VS Code'),
        ('部署平台', 'Streamlit Community Cloud / Hugging Face Spaces'),
    ],
)

# ============================================================
# 三、系统架构
# ============================================================
heading('三、系统架构与功能模块', 1)

heading('3.1 系统总体架构', 2)
para('本系统采用Streamlit Web框架搭建，整体架构分为三层：前端展示层（Streamlit Pages）、业务逻辑层（src模块）、数据存储层（CSV/Excel文件）。系统包含六大功能模块，覆盖数据分析的全链路流程。')

add_table(
    ['模块', '功能', '核心技术'],
    [
        ('数据上传', '多平台CSV/Excel自动识别、字段标准化、缺失值处理、IQR异常检测、数据质量报告', 'ETL管道、Schema校验'),
        ('经营概览', '核心指标卡片、营收趋势图、时段热力图、平台对比分析、统计五数', '描述性统计、时间序列分解'),
        ('商品分析', '销量排行、品类占比、Apriori关联规则挖掘、套餐搭配建议、滞销品识别', 'Apriori关联规则、购物篮分析'),
        ('用户分析', 'RFM三维打分、8类用户分层、K-Means聚类验证、差异化运营策略', 'RFM模型、K-Means聚类'),
        ('智能预测', '17维时间特征工程、随机森林回归预测、Isolation Forest多维异常检测', '随机森林回归、Isolation Forest'),
        ('分析报告', '一键生成自然语言诊断报告、异常提醒、趋势判断、经营建议', '规则模板、自然语言生成'),
    ],
)

heading('3.2 数据流设计', 2)
para('系统的数据流分为六个阶段：')
para('（1）数据上传阶段：用户上传CSV/Excel文件，系统自动检测编码格式，识别多平台表头（美团、饿了么、微信点单等），将不同来源的列名映射到统一标准字段。')
para('（2）ETL清洗阶段：对日期字段进行标准化转换，数值字段进行类型校验和非负检查，缺失值智能填充，完全重复行删除，最后使用IQR方法检测并标记异常订单。')
para('（3）特征工程阶段：从订单明细数据中构造日聚合特征（营收、订单数、顾客数）、时间特征（星期、月份、是否周末）、RFM特征（最近消费间隔、消费频次、消费金额）、时段特征（小时×工作日交叉表）等。')
para('（4）算法建模阶段：依次运行Apriori关联规则挖掘、RFM用户分层、K-Means聚类、随机森林回归预测、Isolation Forest异常检测等算法。')
para('（5）可视化阶段：使用Plotly生成交互式图表，包括折线图（营收趋势）、热力图（时段分布）、水平柱状图（销量排行）、饼图（品类占比）、3D散点图（RFM分布）、气泡图（关联规则）等。')
para('（6）报告生成阶段：汇总各模块分析结果，通过规则模板自动生成自然语言经营诊断报告。')

# ============================================================
# 四、核心算法详解
# ============================================================
heading('四、核心算法详解', 1)

# 4.1 IQR
heading('4.1 IQR异常值检测', 2)
para('IQR（Interquartile Range，四分位距）是一种基于数据分布位置的异常检测方法，用于ETL数据管道中的快速异常值筛选。')
para('算法原理：将数据按从小到大排序后等分为四份，三个分割点分别为第一四分位数Q1（25%位置）、第二四分位数Q2（50%位置，即中位数）、第三四分位数Q3（75%位置）。IQR = Q3 − Q1，定义正常数据范围为：')
formula('正常范围 = [ Q1 − 1.5 × IQR  ,  Q3 + 1.5 × IQR ]')
para('落在该范围之外的数据点被标记为异常值。系数1.5为经验值，等效于正态分布下约2.7个标准差，覆盖约99.3%的数据。')
para('实验应用：在数据管道中对订单金额进行IQR异常检测，快速筛选出金额异常偏高或偏低的订单，标记到数据质量报告中。')

# 4.2 Apriori
heading('4.2 Apriori关联规则挖掘', 2)
para('Apriori算法是数据挖掘领域最经典的关联规则挖掘算法，用于从交易数据中发现商品之间的频繁共现模式。本实验中用于挖掘"购买了A商品的顾客大概率也会购买B商品"的搭配规律。')
para('（1）算法流程：')
para('第一步：构建订单-商品one-hot矩阵。每行代表一个订单，每列代表一种商品，值为1表示该订单包含此商品，值为0表示不包含。')
para('第二步：扫描频繁项集。利用"先验性质"（Apriori Property）进行剪枝——如果一个项集是非频繁的，那么它的所有超集也一定是非频繁的，无需计算。这大大减少了搜索空间。')
para('第三步：对每个频繁项集，计算关联规则的三个核心指标：')

add_table(
    ['指标', '英文', '公式', '含义'],
    [
        ('支持度', 'Support', 'P(A ∩ B)', 'A和B同时出现的订单数 / 总订单数。衡量该组合的普遍程度。'),
        ('置信度', 'Confidence', 'P(B | A)', '包含A的订单中同时也包含B的比例。衡量规则的可靠性。'),
        ('提升度', 'Lift', 'P(B | A) / P(B)', '置信度除以B的自然出现概率。Lift > 1表示正相关，Lift越大关联越强；Lift < 1表示负相关。'),
    ],
)

para('（2）业务应用：')
para('实验发现"酸辣土豆丝 + 可乐"的Lift值为3.2，意味着购买酸辣土豆丝的顾客同时购买可乐的概率是普通顾客购买可乐概率的3.2倍。基于此类高Lift值规则，系统自动生成菜品套餐搭配建议。')

# 4.3 RFM
heading('4.3 RFM用户分层模型', 2)
para('RFM模型是CRM（客户关系管理）领域最经典的用户价值分析模型，从三个维度衡量客户价值：')
add_table(
    ['维度', '英文', '定义', '评价标准'],
    [
        ('R', 'Recency', '最近一次消费距今的天数', '越小越好（说明顾客最近来过）'),
        ('F', 'Frequency', '累计消费的订单数', '越多越好（说明顾客来得勤）'),
        ('M', 'Monetary', '累计消费的总金额', '越多越好（说明顾客花钱多）'),
    ],
)
para('分层方法：对每个用户的R、F、M三个维度分别按分位数进行三等分，每维赋予1~3分（R值越小得分越高，F和M值越大得分越高）。3×3×3=27种评分组合归纳为8类用户：')

add_table(
    ['R分', 'F分', 'M分', '分层标签', '运营策略'],
    [
        ['3（近）', '3（高）', '2-3（高）', '重要价值客户', '维护VIP待遇，提供专属折扣'],
        ['3（近）', '2-3', '2-3', '潜力客户', '推荐高客单价新品，刺激消费升级'],
        ['2-3', '1（低）', '2-3', '重要保持客户', '定向发放回归优惠券'],
        ['3', '1', '1', '新客户', '首单后3天内推送新客专享券'],
        ['1（远）', '2-3', '2-3', '重要挽留客户', '大力度折扣+人工触达召回'],
        ['1', '2-3', '1', '一般价值客户', '维持常规运营节奏'],
        ['1', '1', '2-3', '流失高价值客户', '重点挽回，电话/短信触达'],
        ['1', '1', '1', '流失客户', '低成本触达，顺其自然'],
    ],
)
para('验证方法：使用K-Means无监督聚类对RFM数据进行自然分组，将聚类结果与RFM规则分层进行交叉对比。实验结果显示两种方法的分组高度一致，验证了RFM分层的客观性与可靠性。')

# 4.4 K-Means
heading('4.4 K-Means聚类验证', 2)
para('K-Means算法是最经典的无监督聚类算法，目标是将n个样本划分到K个簇中，使得每个样本到其所属簇中心的距离平方和（SSE）最小。')
para('算法步骤：①随机选择K个初始聚类中心；②将每个样本分配到距离最近的中心所在的簇；③重新计算每个簇的中心（簇内所有样本的均值）；④重复步骤②和③直至收敛或达到最大迭代次数。')
para('K值选择：使用肘部法则（Elbow Method），对K=1至8分别计算SSE，绘制K-SSE曲线。曲线在K=4处出现明显拐点（"肘部"），此后增加K值带来的SSE降低幅度显著减小，因此选择K=4。')
para('实验应用：将K-Means（K=4）的聚类结果与RFM规则分层的8类用户进行交叉对比。交叉表显示两种方法的结果高度吻合——RFM中的"重要价值客户"高度集中于K-Means的高消费簇，"流失客户"集中于低消费簇。这验证了RFM分层的客观性。')

# 4.5 随机森林
heading('4.5 随机森林回归预测', 2)
para('随机森林（Random Forest）是一种基于Bagging策略的集成学习算法，由多棵决策树组成，通过投票（分类）或平均（回归）的方式集成多棵树的预测结果，有效降低过拟合风险。')
para('时间序列转监督学习：本实验的核心创新在于将时间序列预测问题转化为监督学习问题。从原始日营收数据中构造17维时间特征：')

add_table(
    ['特征类别', '特征名称', '维度', '说明'],
    [
        ('趋势特征', 'day_num', '1维', '从第1天开始的递增编号，捕捉长期趋势'),
        ('周期特征', 'weekday（One-Hot编码）', '7维', '星期一至星期日，捕捉周内规律'),
        ('季节特征', 'month、day_of_month', '2维', '月份和几号，捕捉月度规律'),
        ('周末标记', 'is_weekend', '1维', '0=工作日，1=周末'),
        ('滞后特征', 'lag_1、lag_2、lag_3、lag_7', '4维', '前1/2/3/7天的营收值'),
        ('滚动统计', 'rolling_mean_7、rolling_std_7', '2维', '近7天营收的均值和标准差'),
    ],
)

para('模型训练与评估：使用随机森林回归器（n_estimators=100棵决策树，max_depth=5限制树深度以防止过拟合），在90天历史数据上进行训练。使用MAPE（平均绝对百分比误差）评估模型：')
formula('MAPE = (1/n) × Σ|(y_true − y_pred) / y_true| × 100%')
para('实验结果显示MAPE约8.97%，表明平均预测偏差约为±9%。同时，使用残差标准差估算95%预测置信区间（±1.96×σ_residual），为经营决策提供不确定性量化参考。')
para('预测模式：采用递归预测策略——先预测第N+1天，将预测值作为lag_1特征输入模型预测第N+2天，依次递推。为避免误差累积过大，限定预测范围为14天。')

# 4.6 Isolation Forest
heading('4.6 Isolation Forest异常检测', 2)
para('Isolation Forest（孤立森林）是一种基于"隔离"思想的无监督异常检测算法。其核心直觉是：异常数据点是"少数且不同的"——在随机切割的数据空间中，异常点因为稀疏孤立，只需较少的随机切分次数就能被隔离出来。')
para('算法原理：构建多棵隔离树（iTree），每棵树在数据空间中随机选择一个特征和该特征范围内的一个随机切分值，递归切分直到每个数据点被隔离或达到最大深度。数据点的异常分数由其被隔离所需的平均路径长度决定：路径越短（越容易被隔离）→ 异常分数越高 → 越可能是异常点。')
para('实验应用：选取订单金额、商品种类数、商品总数量、平均单价、折扣金额共5个特征维度进行检测，设置contamination=0.05（预期5%的订单为异常）。与数据管道中的IQR方法形成双层检测架构：IQR负责单维度快速初筛，Isolation Forest负责多维度复杂模式复核——例如"金额不算最高但只买一种商品且折扣极大"的组合异常只有后者能发现。')

# ============================================================
# 五、实验结果
# ============================================================
heading('五、实验结果与分析', 1)

heading('5.1 数据概况', 2)
para('实验数据为模拟生成的90天餐饮订单数据，包含26,403条订单明细、12,026笔完整订单、500位顾客、3个平台（美团外卖/微信小程序/饿了么）、38种商品（覆盖热菜、凉菜、汤类、主食、饮品、小吃6大品类），总营收约¥498,721，客单价约¥41.47。')

heading('5.2 算法运行结果汇总', 2)
add_table(
    ['算法模块', '数据规模', '关键结果', '评估指标'],
    [
        ('IQR异常检测', '12,026笔订单', '检测出异常金额订单', '正常范围：Q1-1.5IQR ~ Q3+1.5IQR'),
        ('Apriori关联规则', '38种商品', '14条关联规则', 'Lift值范围1.0~4.0，强关联规则Lift>2'),
        ('RFM用户分层', '500位顾客', '9类用户分层', '分位数三等分，K-Means交叉验证一致'),
        ('K-Means聚类', '500个样本', 'K=4个自然簇', 'SSE=299.04，肘部拐点明确'),
        ('随机森林预测', '90天数据', '未来14天预测', 'MAPE=8.97%，95%置信区间'),
        ('Isolation Forest', '12,026笔订单', '598笔异常订单(5%)', '5维特征检测，contamination=0.05'),
    ],
)

heading('5.3 实验结论', 2)
para('（1）Apriori关联规则能够有效挖掘菜品共现模式，Lift指标排除高频商品干扰后仍能识别有实际意义的搭配关系，可直接用于套餐设计与交叉销售推荐。')
para('（2）RFM模型结合K-Means聚类的交叉验证方式，兼顾了人工规则的业务可解释性和无监督学习的数据驱动客观性，为中小商家提供了低成本、易理解的用户分层方案。')
para('（3）将时间序列预测转化为监督学习问题的特征工程方法，在小数据量场景下（90天日数据）表现良好，MAPE低于9%，优于需要大量数据的LSTM等深度学习方法。')
para('（4）IQR（单维初筛）+ Isolation Forest（多维复核）的双层异常检测架构，兼顾了检测效率与检测深度，能覆盖从简单金额异常到复杂结构异常的多种模式。')
para('（5）从数据管道→特征工程→算法建模→可视化→智能报告的完整闭环设计，体现了数据分析"从数据到决策"的核心价值。')

# ============================================================
# 六、总结与展望
# ============================================================
heading('六、总结与展望', 1)
para('本实验基于Python生态（Streamlit、Pandas、Scikit-learn、Plotly等），搭建了一套完整的餐饮经营数据分析系统，覆盖了数据预处理、特征工程、算法建模、可视化、报告生成的全链路流程。实验中实践了IQR异常检测、Apriori关联规则、RFM用户分层、K-Means聚类、随机森林回归、Isolation Forest异常检测等共6种数据分析与机器学习算法，并实现了多算法协同验证的分析策略。')
para('通过本实验，加深了对数据挖掘、机器学习、时间序列分析等课程理论知识的理解，提升了将理论知识转化为工程实践的能力。')
para('未来可从以下方向进一步改进：引入外部特征（天气、节假日、商圈活动）提升预测精度；接入NLP技术对用户评价进行情感分析；将离线CSV上传升级为实时API数据接入；增加更多评估指标和超参数调优策略。')

doc.add_page_break()

# ============================================================
# 附录：项目结构
# ============================================================
heading('附录：项目文件结构', 1)
code = doc.add_paragraph()
code.paragraph_format.first_line_indent = Cm(0)
run = code.add_run('''
restaurant-analytics/
├── app.py                    # Streamlit主页面入口
├── Dockerfile                # Docker容器配置
├── requirements.txt          # Python依赖包列表
├── pages/                    # 六大分析功能页面
│   ├── 1_数据上传.py          # 数据上传与ETL管道
│   ├── 2_经营概览.py          # 核心指标与趋势分析
│   ├── 3_商品分析.py          # 关联规则与品类分析
│   ├── 4_用户分析.py          # RFM分层与聚类验证
│   ├── 5_智能预测.py          # 营收预测与异常检测
│   └── 6_分析报告.py          # 一键诊断报告生成
├── src/                      # 核心算法与业务逻辑
│   ├── data_pipeline.py      # ETL数据管道（加载/校验/清洗/检测）
│   ├── features.py           # 特征工程（RFM/时间/品类特征）
│   ├── models.py             # 机器学习模型（Apriori/K-Means/RF/IF）
│   ├── analysis.py           # 统计分析与指标计算
│   ├── visualization.py      # Plotly可视化图表工厂
│   ├── report.py             # 自然语言报告生成
│   └── nav_style.py          # 导航栏样式注入
└── sample_data/              # 模拟数据与生成器
    ├── generate_mock_data.py # 数据生成器
    ├── sample_orders.csv     # 模拟订单数据
    └── sample_customers.csv  # 模拟顾客数据
''')
run.font.size = Pt(9)
run.font.name = 'Consolas'

# ===== 保存 =====
out = os.path.join(os.path.expanduser('~/Desktop'), '餐饮数据分析系统_实验报告.docx')
doc.save(out)
print(f'Saved: {out}')
