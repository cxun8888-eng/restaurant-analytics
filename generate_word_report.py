"""生成面试手册 — 对齐简历5条核心职责"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

doc = Document()

for section in doc.sections:
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

style = doc.styles['Normal']
style.font.name = 'Arial'
style.font.size = Pt(10.5)


def h(text, level=2):
    doc.add_heading(text, level=level)


def p(text, bold=False):
    r = doc.add_paragraph()
    run = r.add_run(text)
    run.bold = bold


def q(text):
    para = doc.add_paragraph()
    para.paragraph_format.left_indent = Cm(0.8)
    run = para.add_run(text)
    run.font.color.rgb = RGBColor(0x4F, 0x46, 0xE5)
    run.font.size = Pt(10)


def table(headers, rows):
    tbl = doc.add_table(rows=1 + len(rows), cols=len(headers))
    tbl.style = 'Light Grid Accent 1'
    for i, hd in enumerate(headers):
        tbl.rows[0].cells[i].text = hd
        for r in tbl.rows[0].cells[i].paragraphs:
            for run in r.runs:
                run.bold = True
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            tbl.rows[ri + 1].cells[ci].text = str(val)
    p('')


# ========== 封面 ==========
p('')
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run('餐饮数据分析系统')
run.bold = True
run.font.size = Pt(20)

sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
sub.add_run('面试准备手册').font.size = Pt(14)

p('')

link = doc.add_paragraph()
link.alignment = WD_ALIGN_PARAGRAPH.CENTER
link.add_run('restaurant-analytics-8fgygzvdk9b9edpxe4rpdc.streamlit.app').font.size = Pt(9)

p('')
p('本手册按简历中5条核心职责逐条展开，每条配「面试怎么讲」和「可能追问」。')

# ==========================================
h('职责1：平台搭建与公网部署', 1)
p('简历原文：', bold=True)
q('使用 Python + Streamlit + Pandas + Scikit-learn 完成平台搭建与公网部署')

p('')
p('面试怎么讲：', bold=True)
q('项目用 Streamlit 做 Web 框架，它是纯 Python 的，不需要写前端代码。分析逻辑写在 src 目录里，和数据展示分离。最后部署到 Streamlit Community Cloud，代码推 GitHub 之后自动部署，得到一个公网可以访问的网址。')

p('')
p('可能追问：', bold=True)
p('为什么用 Streamlit 而不是 React/Vue？')
q('项目重点是数据分析能力，不是前端工程。Streamlit 让我把精力放在算法上。需要复杂交互的话，后端分析逻辑可以复用。')

p('项目结构是怎样的？')
q('pages 目录放6个页面，src 目录放核心逻辑（数据管道、特征工程、模型、分析、可视化、报告），和页面解耦，方便测试和维护。')

# ==========================================
h('职责2：数据处理流程', 1)
p('简历原文：', bold=True)
q('设计数据处理流程，支持多平台 CSV/Excel 导入、缺失值处理及异常数据检测')

p('')
p('面试怎么讲：', bold=True)
q('商家从美团、饿了么、微信点单导出的 CSV 表头都不一样。我写了一个数据管道，自动识别不同平台的列名（比如"订单编号"和"order_id"映射到同一个标准字段），然后做缺失值填充、重复删除，最后用 IQR 方法检测异常值——超出 Q1-1.5×IQR 到 Q3+1.5×IQR 范围的标记为异常。')

p('')
p('可能追问：', bold=True)
p('IQR 是什么？')
q('四分位距。Q1 是第25百分位，Q3 是第75百分位，IQR=Q3-Q1。正常范围是 Q1-1.5IQR ~ Q3+1.5IQR。1.5是经验值，等效正态分布下约2.7个标准差。')

p('为什么不用 Z-score？')
q('Z-score 假设数据服从正态分布，但订单金额通常是右偏的。IQR 基于分位数，不受偏态影响，更稳健。')

# ==========================================
h('职责3：用户分层与关联规则', 1)
p('简历原文：', bold=True)
q('应用 RFM 模型进行用户分层，利用 Apriori 关联规则挖掘菜品搭配规律')

p('')
p('面试怎么讲：', bold=True)
q('RFM 是 CRM 领域很经典的用户分层模型。我按最近消费时间（R）、消费频次（F）、消费金额（M）三个维度，用分位数给每个用户打 1-3 分，组合成 8 类——比如高分的是"重要价值客户"，R低但F和M高的属于"重要挽留客户"。')

q('关联规则用 Apriori 算法，找"买了A的顾客大概率也买B"的组合。用三个指标衡量：支持度（普遍不普遍）、置信度（靠不靠谱）、提升度（相关性有多强）——提升度大于1说明正相关，大于2就是强关联，可以建议商家打包成套餐。')

p('')
p('可能追问：', bold=True)
p('三个指标哪个最重要？')
q('提升度Lift。支持度只看出现频率，置信度会被高频商品误导。Lift 排除了这个偏差，衡量的是真正的相关性。')

p('8类怎么分的？')
q('R/F/M 各打 1-3 分，按组合判断。比如 R高(最近来过)+F高+M高 = 重要价值客户；R低(很久没来)+F高+M高 = 重要挽留客户，应该发优惠券召回。')

# ==========================================
h('职责4：营收预测与异常识别', 1)
p('简历原文：', bold=True)
q('基于随机森林回归进行营收预测，并结合 Isolation Forest 实现异常订单识别')

p('')
p('面试怎么讲：', bold=True)
q('营收预测我用的随机森林回归。时间序列本身不能直接喂给机器学习模型，我做了特征工程——构造了星期几、是否周末、前几天的营收（滞后特征）、近7天滚动均值等特征，把时序问题转成了监督学习问题，然后训练随机森林。')

q('异常检测用的 Isolation Forest。思路是随机切割数据空间，异常点因为稀疏孤立，切几下就分出来了。选了订单金额、商品种类数、商品数量、均价、折扣金额5个维度，比单看金额更全面。')

p('')
p('可能追问：', bold=True)
p('为什么用随机森林而不是 LSTM？')
q('LSTM 需要大量数据，90天的日数据不够。随机森林数据量要求不高，而且特征工程做好了效果不差。更重要的是可解释——特征重要性能告诉你是哪些因素在驱动预测。')

p('为什么用 Isolation Forest 而不是只看金额？')
q('只看金额只能抓到"大单异常"。但有些异常是多维的——金额不高但只买了一个商品且折扣很大，可能是刷单。Isolation Forest 在多维特征空间里能发现这种组合异常。')

# ==========================================
h('职责5：可视化与诊断报告', 1)
p('简历原文：', bold=True)
q('集成 Plotly 可视化与自动诊断报告，完成从数据处理到经营建议的分析闭环')

p('')
p('面试怎么讲：', bold=True)
q('用 Plotly 做了交互式图表——折线图看营收趋势、热力图看时段分布、3D散点图看用户分群、气泡图看关联规则。最后做了一个报告生成模块，把前面的分析结果汇总成自然语言报告，告诉商家"哪几款菜建议打包套餐""哪些用户建议发券召回"。')

q('我觉得数据分析的最终价值不是出图表，而是输出可执行的决策。这个报告模块就是完成这个闭环——从原始 CSV，到中间各种算法分析，最后变成商家能直接用的经营建议。')

p('')
p('可能追问：', bold=True)
p('报告是怎么生成的？')
q('基于分析结果写了一套规则模板。比如关联规则 Lift>2 的输出"建议打包套餐"，RFM 中"重要挽留"类用户超过一定数量就输出"建议定向召回"。数据驱动内容，模板控制格式。')

# ==========================================
doc.add_page_break()
h('通用问题', 1)

p('说说这个项目')
q('商家从美团、饿了么、微信点单导出订单 CSV，上传后自动完成清洗、分析、建模、可视化，最后输出经营建议。做了5个分析模块：数据上传、经营概览、商品分析（含关联规则）、用户分析（含RFM分层）、营收预测与异常检测，最后有一键分析报告。用 Streamlit 做的，已部署到公网。')

p('')
p('数据是真实的吗？')
q('用的是模拟数据，但数据生成器模拟了90天、日均约120单、500位顾客、3个平台。内置了真实的业务规律——周末效应、午餐晚餐高峰、商品之间的关联购买模式。分析结果和真实数据表现一致。更换真实数据的话，分析管道不需要改动。')

p('')
p('技术难点')
q('一个是时间序列怎么变成机器学习问题——构造了滞后值、滚动统计等特征。另一个是怎么让多个算法协同而不是孤立——RFM 人工分层用 K-Means 聚类验证、IQR 初筛加 Isolation Forest 复核，互相印证。')

p('')
p('如果数据量更大，会怎么改进？')
q('一是加外部特征——天气、节假日日历，预测会更准。二是可以加 NLP 做评价情感分析。三是把离线上传升级为 API 实时接入。')

# ==========================================
doc.add_page_break()
h('速查表（面试前看这一页就够了）', 1)

table(
    ['职责', '算法', '一句话', '关键数字'],
    [
        ['数据处理', 'IQR', '四分位距检测异常，Q1-1.5IQR~Q3+1.5IQR', '等效2.7个标准差'],
        ['用户分层', 'RFM', 'R近度/F频次/M金额 打分→8类', '分位数三等分'],
        ['关联规则', 'Apriori', '买了A也买B，三个指标衡量', 'Lift>1正相关'],
        ['用户验证', 'K-Means', '无监督聚类，交叉验证RFM', '肘部法则选K'],
        ['营收预测', '随机森林', '时序→监督学习，特征工程+回归', '17维特征'],
        ['异常检测', 'IsolationForest', '随机切割隔离异常点，多维检测', '5维特征'],
    ],
)

p('面试万能结尾（说完技术之后补一句）：')
q('"数据分析的最终价值不是出图表，而是输出可执行的决策。这个项目从原始 CSV 到最后的经营建议报告，完成了一个完整闭环。"')

# ===== 保存 =====
out = os.path.join(os.path.expanduser('~/Desktop'), '餐饮数据分析_面试手册.docx')
doc.save(out)
print(f'Saved: {out}')
